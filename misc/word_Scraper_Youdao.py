# %%
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
import argparse

# Command Parameter
parser = argparse.ArgumentParser()
parser.add_argument('-w', '--word', default = "test", help = 'input word', type = str)
parser.add_argument('-t', '--type', default = 1, help = 'input word type index in collins', type = int)
parser.add_argument('-i', '--index', default = 1, help = 'index of the meaning', type = int)
parser.add_argument('-f', '--filename', default = 'Words.docx', help = 'Microsoft Word file name')
parser.add_argument('-m', '--meaning', default = '-', help = 'Add the meaning of the word')
parser.add_argument('-e', '--episode', default = 'Friends107', help = 'Show the episode the word shows up')

args = parser.parse_args()

word = args.word
word_type = args.type
collins_idx = args.index
doc_name = args.filename
add_meaning = args.meaning
episode = args.episode

# Scrape
url = "http://dict.youdao.com/w/" + word + "/#keyfrom=dict2.top"
results = requests.get(url)
# print(results.status_code)
soup = BeautifulSoup(results.text, "html.parser")
# print(soup.prettify())

# Find pronounce
pron = soup.find_all('span', class_ = 'pronounce')
pron_us = pron[1].get_text(' ', strip = True) if len(pron) > 1 else '-'
print("\n" + pron_us)

# Find the collins meaning and example
collins = soup.find('div', attrs = {'id': 'collinsResult'})

if collins:
    # If the word has more than one 词性
    collins = collins.select('.wt-container:nth-child(' + str(word_type) + ')')
    # print(collins)

    # collins_item: the meaning we need including meaning and examples 
    collins_item = collins[0].select('li:nth-child(' + str(collins_idx) + ') > .collinsMajorTrans')
    collins_meaning = collins_item[0].get_text(' ', strip = True)

    collins_example_div = collins[0].select('li:nth-child(' + str(collins_idx) + ') .examples')
    collins_example = collins_example_div[0].get_text('\n', strip = True) if collins_example_div else '-'
else:
    collins_meaning = add_meaning
    collins_example = '-'

print ("\n" + collins_meaning)
print ("\n" + collins_example)

# %%

# Double check
yOrN = input('\nContent Checked, Y or N?\n')

if yOrN.upper() == 'Y':
    # Write into Microsoft Word
    doc = Document(doc_name)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(10.5)

    doc.add_paragraph(word + ': ' + pron_us)
    p = doc.add_paragraph()
    p.add_run(collins_meaning).bold = True
    run = p.add_run(' [' + episode + ']')
    run.font.color.rgb = RGBColor(0x4E, 0xAD, 0xEA)
    run.bold = True

    doc.add_paragraph('\nex: ' + collins_example)
    doc.add_paragraph('')

    paragraphs = doc.paragraphs
    for i in (-3, -2):
        paragraphs[i].paragraph_format.left_indent = Cm(1)

    doc.save(doc_name)
    print("Write into " + doc_name + " DONE")

else:
    print("Cancel")
# %%
